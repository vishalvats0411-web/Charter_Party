import os
import fitz  # PyMuPDF
import docx
import requests
import json
import re
import time
import sys
import logging

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Configuration ---
API_KEY = "AIzaSyB5RchuqD6syJmkabHygzZAEFqLKHFCCk8"  # Replace with your actual key

# --- File Reading Functions ---

def read_document(file_path: str) -> str:
    """Reads a document, auto-detecting PDF or DOCX."""
    if not os.path.exists(file_path):
        logger.error(f"File not found at '{file_path}'")
        return ""
    
    _, file_extension = os.path.splitext(file_path)
    try:
        if file_extension.lower() == '.pdf':
            return read_pdf_text(file_path)
        elif file_extension.lower() == '.docx':
            return read_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            return ""
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {e}")
        return ""

def read_pdf_text(file_path: str) -> str:
    """Reads text content from a PDF file."""
    try:
        with fitz.open(file_path) as doc:
            return "".join(page.get_text("text") for page in doc)
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path}: {e}")
        return ""

def read_docx(file_path: str) -> str:
    """Reads text content from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

# --- Gemini API Call ---

def get_replacement_instructions(base_cp_text: str, recap_text: str):
    """Sends Base CP and Recap text to Gemini API and gets JSON replacement instructions."""
    if not API_KEY or API_KEY == "YOUR_API_KEY_HERE":
        logger.error("API_KEY is not set or is the default value.")
        return None
        
    logger.info("Connecting to Gemini API...")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro-latest:generateContent?key={API_KEY}"
    
    json_schema = {
        "type": "OBJECT",
        "properties": {
            "replacements": {
                "type": "ARRAY",
                "items": {
                    "type": "OBJECT",
                    "properties": {
                        "old_text": {"type": "STRING"},
                        "new_text": {"type": "STRING"}
                    },
                    "required": ["old_text", "new_text"]
                }
            }
        }
    }

    system_prompt = (
        "You are an expert maritime paralegal. Compare the 'Base CP' and 'Recap' documents "
        "and create a JSON object listing text from Base CP that should be replaced "
        "with updated text from Recap. Preserve exact multi-line formatting in 'old_text'."
    )
    
    # Truncate text if it's too long (Gemini has token limits)
    max_length = 30000  # Reduced from original to avoid token limits
    if len(base_cp_text) > max_length:
        logger.warning(f"Base CP text too long ({len(base_cp_text)} chars), truncating to {max_length}")
        base_cp_text = base_cp_text[:max_length]
    
    if len(recap_text) > max_length:
        logger.warning(f"Recap text too long ({len(recap_text)} chars), truncating to {max_length}")
        recap_text = recap_text[:max_length]
    
    combined_text = f"--- BASE CHARTER PARTY TEXT ---\n{base_cp_text}\n\n--- RECAP DOCUMENT TEXT ---\n{recap_text}"
    
    payload = {
        "systemInstruction": {"parts": [{"text": system_prompt}]},
        "contents": [{"parts": [{"text": combined_text}]}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "responseSchema": json_schema
        }
    }
    
    headers = {'Content-Type': 'application/json'}
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            logger.info(f"Sending API request (Attempt {attempt+1}/{max_retries})...")
            response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=120)
            response.raise_for_status()
            
            result = response.json()
            logger.info("API response received successfully")
            
            if not result.get('candidates'):
                error_msg = result.get('error', {}).get('message', 'API returned no candidates.')
                logger.error(f"API error: {error_msg}")
                return None
                
            json_string = result['candidates'][0]['content']['parts'][0]['text']
            logger.info("Successfully parsed API response")
            
            return json.loads(json_string)
            
        except requests.exceptions.HTTPError as e:
            logger.error(f"HTTP error: {e}")
            if response.status_code == 429:
                logger.warning("Rate limited, waiting before retry...")
                time.sleep(10)  # Wait longer for rate limits
            elif response.status_code >= 500:
                logger.warning("Server error, waiting before retry...")
                time.sleep(5)
            else:
                # For other HTTP errors, try to get error details
                try:
                    error_details = response.json()
                    logger.error(f"API error details: {error_details}")
                except:
                    logger.error(f"HTTP error: {e}")
                
            if attempt < max_retries - 1:
                wait_time = 2 ** (attempt + 1)
                logger.info(f"Retrying in {wait_time}s...")
                time.sleep(wait_time)
            else:
                logger.error(f"API failed after {max_retries} attempts")
                return None
                
        except requests.exceptions.Timeout:
            logger.error("Request timed out")
            if attempt < max_retries - 1:
                wait_time = 2 ** (attempt + 1)
                logger.info(f"Retrying in {wait_time}s...")
                time.sleep(wait_time)
            else:
                logger.error(f"API timed out after {max_retries} attempts")
                return None
                
        except requests.exceptions.RequestException as e:
            logger.error(f"Request exception: {e}")
            if attempt < max_retries - 1:
                wait_time = 2 ** (attempt + 1)
                logger.info(f"Retrying in {wait_time}s...")
                time.sleep(wait_time)
            else:
                logger.error(f"Request failed after {max_retries} attempts")
                return None
                
        except Exception as e:
            logger.error(f"Unexpected error during API call: {e}")
            if attempt < max_retries - 1:
                wait_time = 2 ** (attempt + 1)
                logger.info(f"Retrying in {wait_time}s...")
                time.sleep(wait_time)
            else:
                logger.error(f"Unexpected error after {max_retries} attempts")
                return None
                
    return None

# --- Advanced Text Replacement ---

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replaces text in a paragraph while preserving formatting."""
    full_text = ''.join([r.text for r in paragraph.runs])
    match_index = full_text.find(old_text)
    if match_index == -1:
        return

    before = full_text[:match_index]
    after = full_text[match_index + len(old_text):]
    new_full_text = before + new_text + after

    # Apply new text across runs
    run_idx = 0
    char_idx = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        run.text = new_full_text[char_idx:char_idx + run_len]
        char_idx += run_len
        run_idx += 1

    if char_idx < len(new_full_text):
        if len(paragraph.runs) > 0:
            paragraph.runs[-1].text += new_full_text[char_idx:]

def edit_word_document_in_place(base_doc_path: str, instructions: dict, output_filename: str):
    """Edits a DOCX file in place using replacement instructions."""
    logger.info("Starting Word document editing...")
    replacements = instructions.get("replacements", [])
    if not replacements:
        logger.warning("No replacement instructions found.")
        return False

    try:
        doc = docx.Document(base_doc_path)
        changes_made = 0

        for item in replacements:
            old_text = item.get("old_text")
            new_text = item.get("new_text")

            if not (old_text and new_text):
                continue

            normalized_old_text = re.sub(r'\s+', ' ', old_text).strip()
            logger.info(f"Replacing: '{normalized_old_text[:40]}...'")

            # Replace in paragraphs
            for para in doc.paragraphs:
                normalized_para = re.sub(r'\s+', ' ', para.text).strip()
                if normalized_old_text in normalized_para:
                    replace_text_in_paragraph(para, old_text, new_text)
                    changes_made += 1

            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        normalized_cell = re.sub(r'\s+', ' ', cell.text).strip()
                        if normalized_old_text in normalized_cell:
                            for para in cell.paragraphs:
                                if old_text in para.text:
                                    replace_text_in_paragraph(para, old_text, new_text)
                                    changes_made += 1

        doc.save(output_filename)
        logger.info(f"Successfully created updated document: '{output_filename}' with {changes_made} changes")
        return True

    except Exception as e:
        logger.error(f"Error editing document: {e}")
        return False

# --- Fallback Processing ---

def fallback_processing(base_cp_path, recap_path, output_path):
    """Fallback processing when API fails - creates a simple merged document."""
    try:
        logger.info("Using fallback processing method")
        
        # Read both documents
        base_text = read_document(base_cp_path)
        recap_text = read_document(recap_path)
        
        # Create a simple merged document
        doc = docx.Document()
        
        # Add title
        doc.add_heading('Generated Charter Party', 0)
        doc.add_paragraph('This document was generated by combining the Base CP and Recap documents.')
        doc.add_paragraph('Note: API processing failed, so this is a simple concatenation.')
        
        # Add section for Base CP
        doc.add_heading('Base Charter Party Content', 1)
        if base_text:
            # Split into paragraphs and add
            for para_text in base_text.split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text)
        else:
            doc.add_paragraph('Could not read Base CP content.')
        
        # Add section for Recap
        doc.add_heading('Recap Content', 1)
        if recap_text:
            # Split into paragraphs and add
            for para_text in recap_text.split('\n'):
                if para_text.strip():
                    doc.add_paragraph(para_text)
        else:
            doc.add_paragraph('Could not read Recap content.')
        
        # Save the document
        doc.save(output_path)
        logger.info(f"Fallback document created: {output_path}")
        
        # Return sample changes for display
        return {
            "replacements": [
                {
                    "old_text": "API processing failed",
                    "new_text": "Used fallback processing method"
                }
            ]
        }
    except Exception as e:
        logger.error(f"Fallback processing also failed: {e}")
        return None

# --- Main Processing Function ---

def process_documents(base_cp_path, recap_path, output_path):
    """Main function to process documents."""
    logger.info("Starting Smart CP Generator")

    base_cp_text = read_document(base_cp_path)
    recap_text = read_document(recap_path)

    if not recap_text or not base_cp_text:
        logger.error("Could not read input files.")
        return False, None

    instructions = get_replacement_instructions(base_cp_text, recap_text)
    
    if instructions:
        logger.info("Replacement instructions generated successfully")
        success = edit_word_document_in_place(base_cp_path, instructions, output_path)
        return success, instructions
    else:
        logger.warning("Failed to get replacement instructions, using fallback method")
        # Use fallback processing
        fallback_result = fallback_processing(base_cp_path, recap_path, output_path)
        return fallback_result is not None, fallback_result

# For standalone execution
if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python smart_cp_generator.py <base_cp_path> <recap_path> <output_path>")
        sys.exit(1)
        
    base_cp_path = sys.argv[1]
    recap_path = sys.argv[2]
    output_path = sys.argv[3]
    
    success, changes = process_documents(base_cp_path, recap_path, output_path)
    if success:
        print(f"Successfully processed documents. Changes: {len(changes.get('replacements', [])) if changes else 0}")
    else:
        print("Failed to process documents")
    
    sys.exit(0 if success else 1)